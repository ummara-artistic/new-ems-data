from django.shortcuts import render, get_object_or_404, redirect
# views.py
import io
import os
import csv
from decimal import Decimal
from datetime import datetime
from calendar import monthrange

from django.shortcuts import get_object_or_404, render
from django.http import HttpResponse
from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from calendar import monthrange
from datetime import datetime, timedelta
from .models import Attendance
from employees.models import Employee


from calendar import monthrange
from datetime import datetime, timedelta
from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import Attendance
from employees.models import Employee

from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from calendar import monthrange
from datetime import datetime, timedelta
from decimal import Decimal
from .models import Attendance
from employees.models import Employee

import csv
from decimal import Decimal
from calendar import monthrange
from datetime import datetime
from django.shortcuts import get_object_or_404, render
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from employees.models import Employee
from attendance.models import Attendance, Allowance, Loan, FestivalDeduction

from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from calendar import monthrange
from datetime import datetime, timedelta
from decimal import Decimal
import csv

from .models import Employee, Attendance, Allowance, Loan, FestivalDeduction

import csv
from decimal import Decimal
from calendar import monthrange
from datetime import datetime
from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from .models import Employee, Attendance, Allowance, Loan, FestivalDeduction

from datetime import datetime
from calendar import monthrange
from decimal import Decimal
from django.contrib.auth.decorators import login_required
from django.shortcuts import get_object_or_404, render
from django.http import HttpResponse
import csv
from .models import Employee, Attendance, Allowance, Loan, FestivalDeduction

from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from calendar import monthrange
from decimal import Decimal
from datetime import datetime
import csv

from .models import Employee, Attendance, Allowance, Loan, FestivalDeduction

from django.shortcuts import render, get_object_or_404
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from datetime import datetime
from calendar import monthrange
from decimal import Decimal
import csv

from employees.models import Employee
from attendance.models import Attendance, Allowance, Loan, FestivalDeduction

from calendar import monthrange
from datetime import datetime, timedelta
from decimal import Decimal, ROUND_HALF_UP
import csv

from django.shortcuts import get_object_or_404, render
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required

from employees.models import Employee
from attendance.models import Attendance, Allowance, Loan, FestivalDeduction


from decimal import Decimal, ROUND_HALF_UP
from calendar import monthrange
from datetime import datetime, timedelta
from django.shortcuts import get_object_or_404, render
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
import csv
from employees.models import Employee
from attendance.models import Attendance, Loan, Allowance, FestivalDeduction

from django.shortcuts import get_object_or_404, render
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from datetime import datetime, timedelta
from calendar import monthrange
from decimal import Decimal, ROUND_HALF_UP
import csv

from employees.models import Employee
from attendance.models import Attendance, Allowance, Loan, FestivalDeduction

from django.shortcuts import get_object_or_404, render
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from datetime import datetime, timedelta
from calendar import monthrange
from decimal import Decimal, ROUND_HALF_UP
import csv

from employees.models import Employee
from attendance.models import Attendance, Allowance, Loan, FestivalDeduction

from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from calendar import monthrange
from datetime import datetime, timedelta
from decimal import Decimal, ROUND_HALF_UP
import json

from .models import Employee, Attendance, Loan, FestivalDeduction, Allowance

from datetime import datetime, timedelta
from calendar import monthrange
from decimal import Decimal, ROUND_HALF_UP
from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
import json

from .models import Employee, Attendance, Loan, FestivalDeduction, Allowance

from decimal import Decimal, ROUND_HALF_UP
from datetime import datetime, timedelta
from calendar import monthrange
from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
import json
from .models import Employee, Attendance, Loan, FestivalDeduction, Allowance
from django.contrib.auth.decorators import login_required

from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .models import Employee, Attendance, Allowance, Loan, FestivalDeduction

from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db.models import Q, Sum, Count
from django.http import JsonResponse
from datetime import datetime, timedelta
from calendar import monthrange
from .models import Attendance

from django.shortcuts import render
from django.db.models import Sum, Count, Q
from .models import Attendance, Employee
from datetime import date

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db.models import Sum
from datetime import datetime, date, timedelta
from calendar import monthrange
from .models import Attendance
from employees.models import Employee

@login_required
def attendance_list(request):
    selected_date = request.GET.get('date')
    if selected_date:
        selected_date = date.fromisoformat(selected_date)
    else:
        selected_date = date.today()

    selected_employee = request.GET.get('employee', '')
    selected_status = request.GET.get('status', '')

    employees = Employee.objects.all()
    attendance_records = []

    for emp in employees:
        attendance_qs = Attendance.objects.filter(employee=emp, date=selected_date)
        if selected_status:
            attendance_qs = attendance_qs.filter(status=selected_status)
        attendance = attendance_qs.first()

        # Default values
        status_display = 'Not Marked'
        late_flag = False
        late_minutes = 0
        worked_hours = 0

        if attendance:
            status_display = attendance.get_status_display()

            # Calculate late arrival
            if attendance.in_time:
                shift_start = datetime.combine(selected_date, datetime.strptime("09:15", "%H:%M").time())
                actual_in = datetime.combine(selected_date, attendance.in_time)
                if actual_in > shift_start:
                    late_flag = True
                    late_minutes = (actual_in - shift_start).seconds // 60

            # ✅ Calculate worked hours with lunch deduction
            if attendance.in_time and attendance.out_time:
                in_dt = datetime.combine(selected_date, attendance.in_time)
                out_dt = datetime.combine(selected_date, attendance.out_time)

                total_worked = (out_dt - in_dt).total_seconds() / 3600  # hours

                # Deduct lunch break: 0.5 hr Mon–Thu, 1 hr Friday
                if selected_date.weekday() == 4:  # Friday
                    total_worked -= 1
                else:
                    total_worked -= 0.5

                worked_hours = round(total_worked, 2) if total_worked > 0 else 0

        attendance_records.append({
            'employee': emp,
            'status': status_display,
            'in_time': attendance.in_time if attendance else None,
            'out_time': attendance.out_time if attendance else None,
            'worked_hours': worked_hours,   # ✅ added worked hours
            'overtime_hours': attendance.overtime_hours if attendance else 0,
            'notes': attendance.notes if attendance else '',
            'is_late': late_flag,
            'late_minutes': late_minutes,
            'attendance_obj': attendance,  # for edit/delete links
        })

    # Aggregate totals
    total_present = Attendance.objects.filter(date=selected_date, status='present').count()
    total_absent = Attendance.objects.filter(date=selected_date, status='absent').count()
    total_half_day = Attendance.objects.filter(date=selected_date, status='half_day').count()
    total_marked = total_present + total_absent + total_half_day
    total_late = Attendance.objects.filter(date=selected_date).exclude(in_time__isnull=True).count()
    total_overtime = Attendance.objects.filter(date=selected_date).aggregate(total=Sum('overtime_hours'))['total'] or 0

    context = {
        'employees': employees,
        'attendance_records': attendance_records,
        'selected_date': selected_date,
        'selected_employee': selected_employee,
        'selected_status': selected_status,
        'total_present': total_present,
        'total_absent': total_absent,
        'total_half_day': total_half_day,
        'total_marked': total_marked,
        'total_late': total_late,
        'total_overtime': total_overtime,
    }
    return render(request, 'attendance/list.html', context)



@login_required
def mark_attendance(request):
    if request.method == 'POST':
        employee_id = request.POST.get('employee')
        date_str = request.POST.get('date')  # this comes as string
        status = request.POST.get('status')
        in_time = request.POST.get('in_time')
        out_time = request.POST.get('out_time')
        overtime_hours = request.POST.get('overtime_hours', 0)
        notes = request.POST.get('notes', '')
        
        employee = get_object_or_404(Employee, id=employee_id)

        # ✅ convert date string to Python date
        try:
            attendance_date = datetime.strptime(date_str, "%Y-%m-%d").date()
        except Exception:
            attendance_date = datetime.now().date()
        
        # Create or update attendance record
        attendance, created = Attendance.objects.get_or_create(
            employee=employee,
            date=attendance_date,   # ✅ correct type
            defaults={
                'status': status,
                'in_time': in_time if in_time else None,
                'out_time': out_time if out_time else None,
                'overtime_hours': float(overtime_hours) if overtime_hours else 0,
                'notes': notes,
            }
        )
        
        if not created:
            attendance.status = status
            attendance.in_time = in_time if in_time else None
            attendance.out_time = out_time if out_time else None
            attendance.overtime_hours = float(overtime_hours) if overtime_hours else 0
            attendance.notes = notes
            attendance.save()
        
        messages.success(request, f'Attendance marked for {employee.full_name}')
        return redirect('attendance:list')
    
    employees = Employee.objects.filter(is_active=True).order_by('name')
    today = datetime.now().date()
    
    context = {
        'employees': employees,
        'today': today,
    }
    
    return render(request, 'attendance/mark.html', context)


@login_required
def edit_attendance(request, pk):
    attendance = get_object_or_404(Attendance, pk=pk)
    
    if request.method == 'POST':
        attendance.status = request.POST.get('status')
        attendance.in_time = request.POST.get('in_time') if request.POST.get('in_time') else None
        attendance.out_time = request.POST.get('out_time') if request.POST.get('out_time') else None
        attendance.overtime_hours = float(request.POST.get('overtime_hours', 0))
        attendance.notes = request.POST.get('notes', '')
        attendance.save()
        
        messages.success(request, 'Attendance updated successfully!')
        return redirect('attendance:list')
    
    return render(request, 'attendance/edit.html', {'attendance': attendance})

@login_required
def delete_attendance(request, pk):
    attendance = get_object_or_404(Attendance, pk=pk)
    
    if request.method == 'POST':
        attendance.delete()
        messages.success(request, 'Attendance record deleted successfully!')
        return redirect('attendance:list')
    
    return render(request, 'attendance/delete.html', {'attendance': attendance})



from calendar import monthrange
from decimal import Decimal, ROUND_HALF_UP
from datetime import datetime, timedelta
from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import Employee, Attendance, Loan, FestivalDeduction, Allowance


from calendar import monthrange
from decimal import Decimal, ROUND_HALF_UP
from datetime import datetime, timedelta
from django.shortcuts import render, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import Employee, Attendance, Loan, FestivalDeduction, Allowance
from django.utils import timezone
import pytz


@login_required
from decimal import Decimal, ROUND_HALF_UP
from datetime import datetime, timedelta
from calendar import monthrange
import pytz
from django.shortcuts import render, get_object_or_404
from django.utils import timezone
from django.contrib.auth.decorators import login_required
from .models import Employee, Attendance, Loan, FestivalDeduction, Allowance


@login_required
def monthly_report(request, employee_id, year=None, month=None):
    employee = get_object_or_404(Employee, id=employee_id)
    tz = pytz.timezone('Asia/Karachi')
    today = timezone.localtime(timezone.now(), tz).date()

    # === Determine month/year ===
    year = int(request.GET.get('year', year if year else today.year))
    month = int(request.GET.get('month', month if month else today.month))
    days_in_month = monthrange(year, month)[1]

    # === Handle POST updates ===
    if request.method == 'POST':
        loan_val = request.POST.get('loan_deduction')
        festival_val = request.POST.get('festival_deduction')
        allowance_val = request.POST.get('allowance')

        if loan_val:
            Loan.objects.update_or_create(
                employee=employee, month=month, year=year,
                defaults={'amount': Decimal(loan_val)}
            )
        if festival_val:
            FestivalDeduction.objects.update_or_create(
                employee=employee, month=month, year=year,
                defaults={'deduction_amount': Decimal(festival_val)}
            )
        if allowance_val:
            Allowance.objects.update_or_create(
                employee=employee, month=month, year=year,
                defaults={'amount': Decimal(allowance_val)}
            )

    # === Salary setup ===
    monthly_salary = Decimal(employee.salary)
    per_day_wage = (monthly_salary / Decimal('30')).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)
    per_hour_wage = (per_day_wage / Decimal('8')).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

    # === Attendance records ===
    attendance_records = Attendance.objects.filter(
        employee=employee, date__year=year, date__month=month
    ).order_by('date')
    attendance_dict = {record.date.day: record for record in attendance_records}

    # === Totals ===
    total_allowance_amount = sum(Decimal(a.amount) for a in Allowance.objects.filter(employee=employee, month=month, year=year))
    total_loan_amount = sum(Decimal(l.amount) for l in Loan.objects.filter(employee=employee, month=month, year=year))
    total_festival_amount = sum(Decimal(f.deduction_amount) for f in FestivalDeduction.objects.filter(employee=employee, month=month, year=year))

    calendar_data = []
    total_present = total_absent = total_half_day = total_late_days = 0
    total_worked_hours = total_overtime_hours = Decimal('0.0')
    total_half_day_deduction = Decimal('0.0')

    # === Main loop ===
    for day in range(1, days_in_month + 1):
        date_obj = datetime(year, month, day).date()
        attendance = attendance_dict.get(day)

        is_sunday = date_obj.weekday() == 6
        worked_hours = overtime_hours = salary_deduction = half_day_deduction = Decimal('0.0')
        late_minutes = 0
        is_late = False
        day_earned = Decimal('0.0')

        # Shift timings
        shift_start = datetime.combine(date_obj, datetime.strptime("09:00", "%H:%M").time())
        shift_end = datetime.combine(date_obj, datetime.strptime("17:30", "%H:%M").time())
        lunch_break = timedelta(minutes=30)

        # Friday shorter shift
        if date_obj.weekday() == 4:
            shift_end = datetime.combine(date_obj, datetime.strptime("17:00", "%H:%M").time())
            lunch_break = timedelta(hours=1)

        # === Attendance logic ===
        if attendance and attendance.status in ['present', 'half_day']:
            if attendance.status == 'present':
                total_present += 1
            else:
                total_half_day += 1

            if attendance.in_time and attendance.out_time:
                actual_in = datetime.combine(date_obj, attendance.in_time)
                actual_out = datetime.combine(date_obj, attendance.out_time)

                # Calculate total worked hours (minus lunch)
                worked_delta = (actual_out - actual_in) - lunch_break
                worked_hours = Decimal(worked_delta.total_seconds() / 3600).quantize(Decimal('0.01'))

                # Late arrival check (after 9:15 AM)
                late_threshold = datetime.combine(date_obj, datetime.strptime("09:15", "%H:%M").time())
                if actual_in > late_threshold:
                    is_late = True
                    late_minutes = (actual_in - shift_start).seconds // 60
                    total_late_days += 1

                # === Half-day & pay logic ===
                if attendance.status == 'half_day':
                    # Half-day fixed (4 hours)
                    half_day_deduction = (per_day_wage / Decimal('2')).quantize(Decimal('0.01'))
                    salary_deduction = half_day_deduction
                    day_earned = (per_day_wage / Decimal('2')).quantize(Decimal('0.01'))
                else:
                    # Present full-day logic
                    if worked_hours < Decimal('4.0'):
                        # Less than 4 hrs => half-day deduction
                        half_day_deduction = (per_day_wage / Decimal('2')).quantize(Decimal('0.01'))
                        salary_deduction = half_day_deduction
                        day_earned = (per_day_wage / Decimal('2')).quantize(Decimal('0.01'))
                    elif worked_hours >= Decimal('8.0'):
                        # 8+ hrs (e.g. 9AM-6PM = 8.5hrs) => full day wage
                        day_earned = per_day_wage
                    else:
                        # Partial (between 4–8 hrs) — proportional pay
                        day_earned = (worked_hours * per_hour_wage).quantize(Decimal('0.01'))

                total_half_day_deduction += half_day_deduction

                # === Overtime ===
                scheduled_hours = Decimal(((shift_end - shift_start - lunch_break).total_seconds()) / 3600)
                if worked_hours > scheduled_hours:
                    overtime_hours = (worked_hours - scheduled_hours).quantize(Decimal('0.01'))
                    overtime_pay = (overtime_hours * per_hour_wage * Decimal('1.5')).quantize(Decimal('0.01'))
                    day_earned += overtime_pay
                else:
                    overtime_pay = Decimal('0.0')

                total_worked_hours += worked_hours
                total_overtime_hours += overtime_hours

        else:
            # Absent
            total_absent += 1
            day_earned = Decimal('0.0')

        # === Calendar status ===
        if is_sunday:
            status_for_calendar = 'holiday'
        elif attendance:
            status_for_calendar = attendance.status
        else:
            status_for_calendar = 'absent'

        calendar_data.append({
            'day': day,
            'date': date_obj,
            'attendance': attendance,
            'is_late': is_late,
            'late_minutes': late_minutes,
            'worked_hours': float(worked_hours),
            'overtime_hours': float(overtime_hours),
            'half_day_deduction': float(half_day_deduction),
            'salary_deduction': float(salary_deduction),
            'total_earned': float(day_earned),
            'is_holiday': is_sunday,
            'status': status_for_calendar
        })

    # === Final totals ===
    earned_salary_amount = sum(Decimal(d['total_earned']) for d in calendar_data)
    deductions_amount = total_loan_amount + total_festival_amount + total_half_day_deduction
    final_salary_amount = (earned_salary_amount + total_allowance_amount - deductions_amount).quantize(Decimal('0.01'))

    # === Month navigation ===
    prev_month, prev_year = (month - 1, year) if month > 1 else (12, year - 1)
    next_month, next_year = (month + 1, year) if month < 12 else (1, year + 1)

    context = {
        'employee': employee,
        'year': year,
        'month': month,
        'month_name': datetime(year, month, 1).strftime('%B'),
        'calendar_data': calendar_data,
        'total_present': total_present,
        'total_absent': total_absent,
        'total_half_day': total_half_day,
        'total_late_days': total_late_days,
        'total_worked_hours': float(total_worked_hours),
        'total_overtime_hours': float(total_overtime_hours),
        'per_day_wage': float(per_day_wage),
        'per_hour_wage': float(per_hour_wage),
        'total_allowance': float(total_allowance_amount),
        'total_loan_deduction': float(total_loan_amount),
        'total_festival_deduction': float(total_festival_amount),
        'total_half_day_deduction': float(total_half_day_deduction),
        'earned_salary_amount': float(earned_salary_amount),
        'deductions_amount': float(deductions_amount),
        'final_salary_amount': float(final_salary_amount),
        'total_salary': float(final_salary_amount),
        'prev_year': prev_year,
        'prev_month': prev_month,
        'next_year': next_year,
        'next_month': next_month,
    }

    return render(request, 'attendance/monthly_report.html', context)




# -------------------------------
# SalarySlipDOCX Class
# -------------------------------
# place at top of file (imports assumed present)
import io
import os
from datetime import datetime, timedelta
from calendar import monthrange
from decimal import Decimal, ROUND_HALF_UP
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from .models import Employee, Attendance, Allowance, Loan, FestivalDeduction


# -------------------------------
# -------------------------------
# SalarySlipDOCX Class (uses precomputed salary record)
# -------------------------------
import os
import io
from decimal import Decimal
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from calendar import monthrange
from django.shortcuts import get_object_or_404
from attendance.models import Attendance
from hr.models import Employee, Loan, FestivalDeduction, Allowance


class SalarySlipDOCX:
    def __init__(self, salary_record, logo_path=None):
        self.salary_record = salary_record
        self.logo_path = self._resolve_logo_path(logo_path)

    def _resolve_logo_path(self, logo_path):
        if not logo_path:
            return None
        if os.path.exists(logo_path):
            return logo_path
        base, ext = os.path.splitext(logo_path)
        for alt_ext in [".jpg", ".jpeg", ".png"]:
            candidate = base + alt_ext
            if os.path.exists(candidate):
                return candidate
        return None

    def generate_docx_bytes(self):
        s = self.salary_record
        month_name = datetime(s['year'], s['month'], 1).strftime("%B")

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Nirmala UI"
        style.font.size = Pt(12)

        # Logo (optional)
        if self.logo_path:
            try:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(self.logo_path, width=Inches(1.5))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_paragraph()
            except Exception as e:
                print(f"⚠️ Logo error: {e}")

        # Headings
        doc.add_heading(f"SALARY SLIP – {month_name} {s['year']}", 0)
        doc.add_heading(f"تنخواہ پرچی – {month_name} {s['year']}", 1)

        # Helper for bilingual sections
        def add_section(title_en, title_ur, rows):
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = title_en
            hdr_cells[1].text = title_ur
            for cell in hdr_cells:
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(13)
            for en_label, ur_label, value in rows:
                if isinstance(value, (int, float, Decimal)):
                    value = max(0, value)
                row = table.add_row().cells
                row[0].text = f"{en_label}: {value}"
                row[1].text = f"{ur_label}: {value}"
                p = row[1].paragraphs[0]
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                rtl = OxmlElement('w:bidi')
                rtl.set(qn('w:val'), 'on')
                p._p.get_or_add_pPr().append(rtl)
            doc.add_paragraph()

        # Employee Info
        add_section(
            "Employee Information", "ملازم کی معلومات",
            [
                ("Name", "نام", s['employee'].full_name),
                ("Employee ID", "ملازم نمبر", f"#{s['employee'].id}"),
                ("Type", "قسم", s['employee'].get_employee_type_display()),
                ("Joining Date", "شمولیت کی تاریخ", s['employee'].joining_date.strftime("%d/%m/%Y")),
                ("Age", "عمر", f"{s['employee'].age} years"),
            ]
        )

        # Attendance Summary
        add_section(
            "Attendance Summary", "حاضری کا خلاصہ",
            [
                ("Present Days", "حاضر دن", s['total_present']),
                ("Half Days", "نصف دن", s['total_half_day']),
                ("Absent Days", "غیر حاضر دن", s['total_absent']),
                ("Late Days", "تاخیر والے دن", s.get('total_late_days', 0)),
                ("Overtime Hours", "اضافی گھنٹے", float(s['total_overtime_hours'])),
            ]
        )

        # Salary Breakdown
        add_section(
            "Salary Breakdown", "تنخواہ کی تفصیل",
            [
                ("Earned Salary", "کمائی شدہ تنخواہ", float(s['earned_salary_amount'])),
                ("Overtime Pay", "اوور ٹائم", float(s.get('overtime_amount', 0))),
                ("Allowance", "الاؤنس", float(s['total_allowance'])),
                ("Loan Deductions", "قرض کی ادائیگی", float(s['total_loan_deduction'])),
                ("Festival Deductions", "تہوار کی کٹوتیاں", float(s['total_festival_deduction'])),
                ("Half-Day Deduction", "نصف دن کی کٹوتی", float(s.get('total_half_day_deduction', 0))),
                ("Final Salary", "حتمی تنخواہ", float(s['final_salary_amount'])),
            ]
        )

        # Total Salary Heading
        doc.add_heading(f"TOTAL SALARY / کل تنخواہ: Rs. {float(s['final_salary_amount']):,.2f}", 1)
        doc.add_paragraph(f"Generated on: {datetime.now():%d/%m/%Y %H:%M}")
        doc.add_paragraph(f"تاریخ: {datetime.now():%d/%m/%Y %H:%M}")

        # Save to BytesIO
        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        return docx_io.read()


# -------------------------------
# Monthly Salary DOCX (Uses Existing Monthly Report Data)
# -------------------------------
@login_required
def generate_monthly_salary_docx(request, employee_id, year=None, month=None):
    """
    This version uses the already computed monthly_report data
    — no recalculation of attendance or salary.
    """
    from attendance.views import monthly_report  # Import the same function
    
    # Call the existing monthly_report logic
    response = monthly_report(request, employee_id, year, month)
    context = response.context_data if hasattr(response, "context_data") else response.context  # Extract context safely

    # Prepare salary data
    salary_data = {
        'employee': context['employee'],
        'year': context['year'],
        'month': context['month'],
        'total_present': context['total_present'],
        'total_half_day': context['total_half_day'],
        'total_absent': context['total_absent'],
        'total_late_days': context.get('total_late_days', 0),
        'total_worked_hours': context['total_worked_hours'],
        'total_overtime_hours': context['total_overtime_hours'],
        'earned_salary_amount': context['earned_salary_amount'],
        'overtime_amount': 0,  # already included in earned_salary if logic merged
        'total_allowance': context['total_allowance'],
        'total_loan_deduction': context['total_loan_deduction'],
        'total_festival_deduction': context['total_festival_deduction'],
        'total_half_day_deduction': context['total_half_day_deduction'],
        'final_salary_amount': context['final_salary_amount'],
    }

    logo_path = os.path.join('static', 'images', 'logo.png')
    docx_bytes = SalarySlipDOCX(salary_data, logo_path=logo_path).generate_docx_bytes()

    response = HttpResponse(
        docx_bytes,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{salary_data['employee'].full_name}_salary_{month}_{year}.docx"'
    return response




from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from datetime import datetime, date
from .models import Employee, SalaryPayment, Attendance  # Assuming SalaryPayment stores monthly salary info

from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from datetime import datetime
from .models import Employee, SalaryPayment, Attendance

from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from datetime import datetime
from .models import Employee, SalaryPayment, Attendance

from django.shortcuts import render
from django.contrib.auth.decorators import login_required
from datetime import datetime
from calendar import monthrange
from decimal import Decimal
from .models import Employee, Attendance, Loan, FestivalDeduction, Allowance

@login_required
def monthly_report_all(request, year=None, month=None):
    today = datetime.now().date()
    year = int(request.GET.get('year', year if year else today.year))
    month = int(request.GET.get('month', month if month else today.month))
    days_in_month = monthrange(year, month)[1]

    employees_data = []

    for employee in Employee.objects.all():
        # Monthly salary
        monthly_salary = Decimal(employee.salary)
        working_days = [d for d in range(1, days_in_month+1)
                        if datetime(year, month, d).weekday() != 6]
        total_working_days = len(working_days)
        per_day_wage = monthly_salary / Decimal(total_working_days)
        per_hour_wage = per_day_wage / Decimal('7.5')

        # Attendance records
        attendance_records = Attendance.objects.filter(
            employee=employee,
            date__year=year,
            date__month=month
        ).order_by('date')
        attendance_dict = {record.date.day: record for record in attendance_records}

        # Allowances and deductions
        total_allowance_amount = sum(
            [a.amount for a in Allowance.objects.filter(employee=employee, month=month, year=year)]
        )
        total_loan_amount = sum(
            [l.amount for l in Loan.objects.filter(employee=employee, month=month, year=year)]
        )
        total_festival_amount = sum(
            [f.deduction_amount for f in FestivalDeduction.objects.filter(employee=employee, month=month, year=year)]
        )

        # Attendance calculations
        total_present = total_absent = total_half_day = total_late_days = 0
        total_overtime_hours = total_worked_hours = Decimal('0.0')

        for day in range(1, days_in_month+1):
            date_obj = datetime(year, month, day).date()
            attendance = attendance_dict.get(day)
            is_sunday = date_obj.weekday() == 6
            worked_hours = overtime_hours = salary_deduction = half_day_deduction = Decimal('0.0')
            late_minutes = 0

            shift_start = datetime.combine(date_obj, datetime.strptime("09:00", "%H:%M").time())
            shift_end = datetime.combine(date_obj, datetime.strptime("17:30", "%H:%M").time())
            if date_obj.weekday() == 4:  # Friday
                shift_end = datetime.combine(date_obj, datetime.strptime("17:00", "%H:%M").time())

            if not is_sunday:
                if attendance:
                    if attendance.status == 'present' and attendance.in_time and attendance.out_time:
                        total_present += 1
                        actual_in = datetime.combine(date_obj, attendance.in_time)
                        actual_out = datetime.combine(date_obj, attendance.out_time)
                        worked_delta = actual_out - actual_in
                        worked_hours = Decimal(worked_delta.total_seconds()) / Decimal('3600')

                        # Overtime
                        if actual_out > shift_end:
                            extra_time = actual_out - shift_end
                            overtime_hours = Decimal(extra_time.total_seconds()) / Decimal('3600')

                        # Late
                        late_threshold = datetime.combine(date_obj, datetime.strptime("09:15", "%H:%M").time())
                        if actual_in > late_threshold:
                            late_minutes = (actual_in - shift_start).seconds // 60
                            total_late_days += 1
                            late_hours = Decimal(late_minutes)/Decimal('60')
                            if overtime_hours >= late_hours:
                                overtime_hours -= late_hours
                            else:
                                salary_deduction += (late_hours - overtime_hours) * per_hour_wage
                                overtime_hours = Decimal('0.0')

                        # Half-day
                        if late_minutes >= 60:
                            half_day_deduction = per_day_wage / 2
                            salary_deduction += half_day_deduction
                            total_half_day += 1

                        total_overtime_hours += overtime_hours
                        total_worked_hours += worked_hours
                    else:
                        total_absent += 1
                else:
                    total_absent += 1

        # Attendance allowance
        if total_present == total_working_days and getattr(employee, 'attendance_allowance', 0):
            total_allowance_amount += employee.attendance_allowance

        earned_salary_amount = float(per_day_wage * total_present)
        overtime_amount = float(total_overtime_hours * per_hour_wage)
        deductions_amount = float(total_loan_amount + total_festival_amount)
        final_salary_amount = earned_salary_amount + overtime_amount + float(total_allowance_amount) - deductions_amount

        employees_data.append({
            'employee': employee,
            'total_present': total_present,
            'total_absent': total_absent,
            'total_half_day': total_half_day,
            'total_late_days': total_late_days,
            'total_worked_hours': float(round(total_worked_hours, 2)),
            'total_overtime_hours': float(round(total_overtime_hours, 2)),
            'total_allowance': float(total_allowance_amount),
            'total_loan_deduction': float(total_loan_amount),
            'total_festival_deduction': float(total_festival_amount),
            'final_salary_amount': float(final_salary_amount),
        })

    prev_month, prev_year = (month-1, year) if month > 1 else (12, year-1)
    next_month, next_year = (month+1, year) if month < 12 else (1, year+1)

    context = {
        'year': year,
        'month': month,
        'month_name': datetime(year, month, 1).strftime('%B'),
        'employees_data': employees_data,
        'prev_year': prev_year,
        'prev_month': prev_month,
        'next_year': next_year,
        'next_month': next_month,
    }

    return render(request, 'attendance/list.html', context)
