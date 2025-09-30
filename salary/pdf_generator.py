import io, os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


class SalarySlipDOCX:
    """
    Generate bilingual salary slip in-memory and return DOCX bytes.
    """
    def __init__(self, salary_record, logo_path=None):
        self.salary_record = salary_record
        self.logo_path = self._resolve_logo_path(logo_path)

    def _resolve_logo_path(self, logo_path):
        """
        Ensure the logo path points to a valid image file.
        Tries .jpg and .png if .jfif is not supported.
        """
        if not logo_path:
            return None

        # If file exists, return it
        if os.path.exists(logo_path):
            return logo_path

        # Try alternative extensions
        base, ext = os.path.splitext(logo_path)
        for alt_ext in [".jpg", ".jpeg", ".png"]:
            candidate = base + alt_ext
            if os.path.exists(candidate):
                return candidate

        print(f"⚠️ Logo not found at {logo_path}")
        return None

    def generate_docx_bytes(self):
        s = self.salary_record
        month_name = datetime(s.year, s.month, 1).strftime("%B")

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "Nirmala UI"
        style.font.size = Pt(12)

        # -------------------
        # Add logo (centered)
        # -------------------
        if self.logo_path:
            try:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.add_picture(self.logo_path, width=Inches(1.5))
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_paragraph()
            except Exception as e:
                print(f"⚠️ Logo error: {e}")

        # Headings
        doc.add_heading(f"SALARY SLIP – {month_name} {s.year}", 0)
        doc.add_heading(f"تنخواہ پرچی – {month_name} {s.year}", 1)

        # Helper for bilingual sections
        def add_section(title_en, title_ur, rows):
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = title_en
            hdr_cells[1].text = title_ur
            for cell in hdr_cells:
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(13)
            for en_label, ur_label, value in rows:
                row = table.add_row().cells
                row[0].text = f"{en_label}: {value}"
                row[1].text = f"{ur_label}: {value}"
                p = row[1].paragraphs[0]
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                rtl = OxmlElement('w:bidi')
                rtl.set(qn('w:val'), 'on')
                p._p.get_or_add_pPr().append(rtl)
            doc.add_paragraph()

        # Employee Info
        add_section(
            "Employee Information", "ملازم کی معلومات",
            [
                ("Name", "نام", s.employee.full_name),
                ("Employee ID", "ملازم نمبر", f"#{s.employee.id}"),
                ("Type", "قسم", s.employee.get_employee_type_display()),
                ("Joining Date", "شمولیت کی تاریخ", s.employee.joining_date.strftime("%d/%m/%Y")),
                ("Age", "عمر", f"{s.employee.age} years"),
            ]
        )

        # Attendance Summary
        add_section(
            "Attendance Summary", "حاضری کا خلاصہ",
            [
                ("Present Days", "حاضر دن", s.days_present),
                ("Absent Days", "غیر حاضر دن", s.days_absent),
                ("Overtime Hours", "اضافی گھنٹے", s.overtime_hours),
            ]
        )

        # Salary Breakdown
        add_section(
            "Salary Breakdown", "تنخواہ کی تفصیل",
            [
                ("Basic Salary", "تنخواہ", f"Rs. {s.basic_salary:,.0f}"),
                ("Overtime", "اوور ٹائم", f"Rs. {s.overtime_amount:,.0f}"),
                ("Tea Allowance", "چائے", f"Rs. {s.tea_allowance:,.0f}"),
                ("Advance Taken", "ایڈوانس", f"Rs. {s.advance_taken:,.0f}"),
                ("Other Deductions", "دیگر کٹوتی", f"Rs. {s.deductions:,.0f}"),
                ("Dues", "بقایا", f"Rs. {s.dues:,.0f}"),
            ]
        )

        # Total Salary
        doc.add_heading(f"TOTAL SALARY / کل تنخواہ: Rs. {s.total_salary:,.0f}", 1)
        doc.add_paragraph(f"Generated on: {datetime.now():%d/%m/%Y %H:%M}")
        doc.add_paragraph(f"تاریخ: {datetime.now():%d/%m/%Y %H:%M}")

        # Save to in-memory bytes
        docx_io = io.BytesIO()
        doc.save(docx_io)
        docx_io.seek(0)
        return docx_io.read()


# Wrapper
def generate_salary_pdf(salary_record, logo_path=None):
    return SalarySlipDOCX(salary_record, logo_path=logo_path).generate_docx_bytes()
