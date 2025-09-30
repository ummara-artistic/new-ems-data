import os
import tempfile
from calendar import monthrange
from datetime import datetime
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db.models import Sum
from django.http import HttpResponse

from .models import SalaryRecord

from attendance.models import Attendance, Holiday
from .pdf_generator import generate_salary_pdf

# -----------------------------------------
# SALARY LIST VIEW
# -----------------------------------------
@login_required
def salary_list(request):
    """
    Show all salary records for the selected month/year (and optional employee)
    and include a grand-total row at the end of the table.
    """
    month_filter = request.GET.get('month', datetime.now().month)
    year_filter = request.GET.get('year', datetime.now().year)
    employee_filter = request.GET.get('employee')

    # Safely convert month/year to integers
    try:
        month_filter = int(month_filter)
        year_filter = int(year_filter)
    except (ValueError, TypeError):
        month_filter = datetime.now().month
        year_filter = datetime.now().year

    salary_records = SalaryRecord.objects.filter(month=month_filter, year=year_filter)
    if employee_filter:
        salary_records = salary_records.filter(employee_id=employee_filter)

    totals = salary_records.aggregate(
        total_basic=Sum('basic_salary'),
        total_overtime=Sum('overtime_amount'),
        total_tea=Sum('tea_allowance'),
        total_advance=Sum('advance_taken'),
        total_deductions=Sum('deductions'),
        total_dues=Sum('dues'),
        grand_total=Sum('total_salary'),
    )

    employees = Employee.objects.filter(is_active=True).order_by('name')
    current_year = datetime.now().year
    years = range(current_year - 2, current_year + 3)
    months = [
        (1, 'January'), (2, 'February'), (3, 'March'), (4, 'April'),
        (5, 'May'), (6, 'June'), (7, 'July'), (8, 'August'),
        (9, 'September'), (10, 'October'), (11, 'November'), (12, 'December')
    ]

    context = {
        'salary_records': salary_records,
        'totals': totals,
        'employees': employees,
        'selected_month': month_filter,
        'selected_year': year_filter,
        'selected_employee': employee_filter,
        'years': years,
        'months': months,
    }

    return render(request, 'salary/list.html', context)


# -----------------------------------------
# GENERATE SALARY VIEW
# -----------------------------------------
from calendar import monthrange
from datetime import datetime, time
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db.models import Sum
from django.http import HttpResponse
import os, tempfile

from .models import SalaryRecord
from employees.models import Employee
from attendance.models import Attendance, Holiday
from .pdf_generator import generate_salary_pdf

# -------------------------
# GENERATE SALARY VIEW

from decimal import Decimal, ROUND_HALF_UP
from calendar import monthrange
from datetime import datetime, timedelta
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db.models import Sum
from .models import SalaryRecord
from employees.models import Employee

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from datetime import datetime, timedelta
from calendar import monthrange
from decimal import Decimal, ROUND_HALF_UP
from django.db.models import Sum

from decimal import Decimal, ROUND_HALF_UP
from datetime import datetime, timedelta
from calendar import monthrange
from io import BytesIO
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from docx import Document
from docx.shared import Pt

@login_required
def generate_salary(request, employee_id, year, month):
    employee = get_object_or_404(Employee, id=employee_id)

    # Check if salary exists
    existing_salary = SalaryRecord.objects.filter(employee=employee, month=month, year=year).first()
    if existing_salary:
        messages.warning(request, f'Salary for {employee.full_name} for {month}/{year} already exists!')
        return redirect('salary:edit', existing_salary.pk)

    # Attendance records
    attendance_records = Attendance.objects.filter(employee=employee, date__year=year, date__month=month)
    attendance_dict = {att.date: att for att in attendance_records}

    # Initialize totals
    days_present = half_days = days_absent = leave_days = total_late_days = 0
    total_overtime = Decimal(0)
    late_minutes_total = 0

    # Build calendar data
    calendar_data = []
    for day in range(1, monthrange(year, month)[1] + 1):
        date_obj = datetime(year, month, day).date()
        attendance = attendance_dict.get(date_obj)
        is_late = False
        late_minutes = 0

        if attendance:
            if attendance.status == 'present':
                days_present += 1
            elif attendance.status == 'half_day':
                half_days += 1
            elif attendance.status == 'absent':
                days_absent += 1
            elif attendance.status == 'leave':
                leave_days += 1

            if attendance.overtime_hours:
                total_overtime += Decimal(attendance.overtime_hours)

            if attendance.in_time:
                shift_start = datetime.combine(date_obj, datetime.strptime("09:15", "%H:%M").time())
                actual_in = datetime.combine(date_obj, attendance.in_time)
                if actual_in > shift_start:
                    is_late = True
                    late_minutes = 30
                    total_late_days += 1
                    late_minutes_total += late_minutes

        calendar_data.append({
            'date': date_obj,
            'day': day,
            'attendance': attendance,
            'is_late': is_late,
            'late_minutes': late_minutes,
            'is_today': date_obj == datetime.today().date()
        })

    late_penalty_hours = Decimal(late_minutes_total / 60).quantize(Decimal('0.01'), rounding=ROUND_HALF_UP)

    # Holidays & Sundays
    holidays = Holiday.objects.filter(date__year=year, date__month=month)
    sundays = [datetime(year, month, day) for day in range(1, monthrange(year, month)[1] + 1) if datetime(year, month, day).weekday() == 6]

    # Leave penalty
    sunday_penalty = 0
    for leave in attendance_records.filter(status='leave'):
        if leave.date.weekday() == 0 and (leave.date - timedelta(days=1)).weekday() == 6:
            sunday_penalty += 1

    leave_penalty_days = leave_days + len(holidays) + sunday_penalty
    working_days = monthrange(year, month)[1] - len(sundays)

    # Rates
    daily_rate = Decimal(employee.salary) / Decimal(working_days)
    hourly_rate = daily_rate / Decimal(8)

    # Salary calculations
    basic_salary = daily_rate * Decimal(days_present + 0.5 * half_days)
    extra_overtime = total_overtime / Decimal(2)
    total_hours = total_overtime + extra_overtime
    overtime_amount = total_hours * hourly_rate * Decimal('1.5')
    deductions = daily_rate * Decimal(days_absent + leave_penalty_days) + (late_penalty_hours * hourly_rate)
    tea_allowance = Decimal(1000)
    total_salary = basic_salary + overtime_amount + tea_allowance - deductions

    if request.method == 'POST':
        # Override values from form
        basic_salary = Decimal(request.POST.get('basic_salary', basic_salary))
        overtime_amount = Decimal(request.POST.get('overtime_amount', overtime_amount))
        tea_allowance = Decimal(request.POST.get('tea_allowance', tea_allowance))
        advance_taken = Decimal(request.POST.get('advance_taken', 0))
        deductions = Decimal(request.POST.get('deductions', deductions))
        dues = Decimal(request.POST.get('dues', 0))
        notes = request.POST.get('notes', '')

        total_salary = basic_salary + overtime_amount + tea_allowance - advance_taken - deductions + dues

        # Save SalaryRecord
        salary_record = SalaryRecord.objects.create(
            employee=employee,
            month=month,
            year=year,
            basic_salary=basic_salary,
            overtime_amount=overtime_amount,
            tea_allowance=tea_allowance,
            advance_taken=advance_taken,
            deductions=deductions,
            dues=dues,
            total_salary=total_salary,
            days_present=days_present,
            days_absent=days_absent,
            half_days=half_days,
            overtime_hours=float(total_overtime),
            late_penalty_hours=late_penalty_hours,
            notes=notes
        )

        # -------- Generate DOCX in memory --------
        document = Document()

        # Title
        heading = document.add_heading("Salary Slip", level=0)
        heading.alignment = 1  # center

        # Employee Info
        document.add_paragraph(f"Employee: {salary_record.employee.full_name}", style="Normal")
        document.add_paragraph(f"Month: {salary_record.month}/{salary_record.year}", style="Normal")

        # Salary Breakdown Table
        table = document.add_table(rows=1, cols=2)
        table.style = "Light Grid Accent 1"
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Description"
        hdr_cells[1].text = "Amount"

        rows = [
            ("Basic Salary", salary_record.basic_salary),
            ("Overtime Amount", salary_record.overtime_amount),
            ("Tea Allowance", salary_record.tea_allowance),
            ("Advance Taken", salary_record.advance_taken),
            ("Deductions", salary_record.deductions),
            ("Dues", salary_record.dues),
            ("Total Salary", salary_record.total_salary),
        ]

        for desc, value in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = desc
            row_cells[1].text = f"{value:.2f}"

        # Notes
        if salary_record.notes:
            document.add_paragraph("\nNotes:", style="Heading 2")
            document.add_paragraph(salary_record.notes)

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)

        # Serve DOCX immediately
        filename = f"salary_slip_{employee.name}_{employee.last_name}_{month}_{year}.docx".replace(" ", "_")
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response

    context = {
        'employee': employee,
        'month': month,
        'year': year,
        'month_name': datetime(year, month, 1).strftime('%B'),
        'days_present': days_present,
        'days_absent': days_absent,
        'half_days': half_days,
        'total_overtime': float(total_overtime),
        'extra_overtime': float(extra_overtime),
        'total_hours': float(total_hours),
        'late_penalty_hours': float(late_penalty_hours),
        'total_late_days': total_late_days,
        'calendar_data': calendar_data,
        'leave_penalty_days': leave_penalty_days,
        'basic_salary': float(basic_salary),
        'overtime_amount': float(overtime_amount),
        'tea_allowance': float(tea_allowance),
        'deductions': float(deductions),
        'total_salary': float(total_salary),
        'working_days': working_days,
    }

    return render(request, 'salary/generate.html', context)






# -------------------------
# GENERATE SALARY PDF VIEW
# -------------------------

from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from .models import SalaryRecord


import os
from django.conf import settings

import os
from django.conf import settings
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from .models import SalaryRecord



@login_required
def generate_pdf(request, pk):
    salary_record = get_object_or_404(SalaryRecord, pk=pk)

    try:
        # Correct logo file path
        logo_path = os.path.join(
            settings.STATICFILES_DIRS[0],
            "images",
            "2ec7ca9c-da9f-44f9-94cd-d93012fa7c21.jfif",
        )

        if not os.path.exists(logo_path):
            messages.warning(request, f"⚠️ Logo not found at {logo_path}")

        # Generate DOCX entirely in memory
        docx_bytes = generate_salary_pdf(salary_record, logo_path=logo_path)

        # Serve as downloadable file
        filename = (
            f"salary_slip_{salary_record.employee.name}_"
            f"{salary_record.employee.last_name}_"
            f"{salary_record.month}_{salary_record.year}.docx"
        ).replace(" ", "_")

        response = HttpResponse(
            docx_bytes,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response

    except Exception as e:
        messages.error(request, f"Error generating salary document: {e}")
        return redirect("salary:list")


# -----------------------------------------
# EDIT SALARY VIEW
# -----------------------------------------
@login_required
def edit_salary(request, pk):
    salary_record = get_object_or_404(SalaryRecord, pk=pk)

    if request.method == 'POST':
        # Use 0 as default if the POST field is empty
        salary_record.basic_salary = float(request.POST.get('basic_salary') or 0)
        salary_record.overtime_amount = float(request.POST.get('overtime_amount') or 0)
        salary_record.tea_allowance = float(request.POST.get('tea_allowance') or 0)
        salary_record.advance_taken = float(request.POST.get('advance_taken') or 0)
        salary_record.deductions = float(request.POST.get('deductions') or 0)
        salary_record.dues = float(request.POST.get('dues') or 0)
        salary_record.notes = request.POST.get('notes', '')

        salary_record.total_salary = (
            salary_record.basic_salary +
            salary_record.overtime_amount +
            salary_record.tea_allowance -
            salary_record.advance_taken -
            salary_record.deductions +
            salary_record.dues
        )

        salary_record.save()
        messages.success(request, 'Salary record updated successfully!')
        return redirect('salary:list')

    context = {
        'salary_record': salary_record,
        'month_name': datetime(salary_record.year, salary_record.month, 1).strftime('%B'),
    }
    return render(request, 'salary/edit.html', context)



# -----------------------------------------
# DELETE SALARY VIEW
# -----------------------------------------
@login_required
def delete_salary(request, pk):
    salary_record = get_object_or_404(SalaryRecord, pk=pk)

    if request.method == 'POST':
        salary_record.delete()
        messages.success(request, 'Salary record deleted successfully!')
        return redirect('salary:list')

    return render(request, 'salary/delete.html', {'salary_record': salary_record})


# -----------------------------------------
